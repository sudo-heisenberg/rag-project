"""Document loading and parsing utilities."""

from pathlib import Path
from typing import List, Dict, Any
import logging
from dataclasses import dataclass

# Document parsers
import pypdf
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

from src.utils import extract_metadata

logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document data structure."""
    content: str
    metadata: Dict[str, Any]
    doc_id: str

class DocumentLoader:
    """Load and parse documents from various formats."""

    def __init__(self):
        """Initialize document loader."""
        self.supported_formats = {
            '.txt': self._load_txt,
            '.pdf': self._load_pdf,
            '.md': self._load_markdown,
            '.docx': self._load_docx,
        }

    def load_document(self, file_path: Path) -> Document:
        """
        Load a single document.

        Args:
            file_path: Path to the document

        Returns:
            Document object
        """
        logger.info(f"Loading document: {file_path}")

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")

        # Load content using appropriate parser
        content = self.supported_formats[extension](file_path)

        # Extract metadata
        metadata = extract_metadata(file_path)
        metadata['source'] = str(file_path)

        # Create document ID
        doc_id = self._generate_doc_id(file_path)

        return Document(content=content, metadata=metadata, doc_id=doc_id)

    def load_directory(self, directory_path: Path, recursive: bool = True) -> List[Document]:
        """
        Load all documents from a directory.

        Args:
            directory_path: Path to the directory
            recursive: Whether to search recursively

        Returns:
            List of Document objects
        """
        logger.info(f"Loading documents from directory: {directory_path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    doc = self.load_document(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {str(e)}")

        logger.info(f"Loaded {len(documents)} documents")
        return documents

    def _load_txt(self, file_path: Path) -> str:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file."""
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)

    def _load_markdown(self, file_path: Path) -> str:
        """Load markdown file and convert to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert markdown to HTML then to plain text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()

    def _load_docx(self, file_path: Path) -> str:
        """Load DOCX file."""
        doc = DocxDocument(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID."""
        # Use file name without extension as ID
        return file_path.stem

# Example usage
if __name__ == "__main__":
    loader = DocumentLoader()

    # Load single document
    # doc = loader.load_document(Path("data/raw/sample.pdf"))
    # print(f"Loaded: {doc.doc_id}, Content length: {len(doc.content)}")

    # Load directory
    # docs = loader.load_directory(Path("data/raw"))
    # print(f"Total documents loaded: {len(docs)}")
